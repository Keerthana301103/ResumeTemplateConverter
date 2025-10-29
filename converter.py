import pdfplumber
import docx
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
import os
from dotenv import load_dotenv

# --- 1. SHARED FUNCTIONS (Used by all templates) ---

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY missing.")

genai.configure(api_key=api_key)
# Using gemini-1.5-flash as a stable, available model
model = genai.GenerativeModel('gemini-2.5-flash') 

# Define the brand color (used by T1, T2 uses it inline)
M_RED = RGBColor(204, 31, 32)


def extract_text_from_pdf(file):
    """Extracts text from an uploaded PDF file."""
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file):
    """Extracts text from an uploaded DOCX file."""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def call_gemini_api(prompt):
    """Calls the Gemini API and returns the text response."""
    response = model.generate_content(prompt)
    try:
        return response.text
    except ValueError:
        print("Warning: Gemini response blocked or empty.")
        return "Error: Could not generate content."

def set_table_no_border(table):
    """Helper function to remove all borders from a table. (Used by T1 and T2)"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)


# --- 2. TEMPLATE 1 FUNCTIONS ---
# (This is your first template's logic)

# --- Helpers for Template 1 ---
def add_heading_t1(doc, text, level=1):
    if not text or text.strip().lower() == 'none': return
    style = f'Heading {level}' if level > 0 else 'Title'
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    p.runs[0].font.color.rgb = M_RED
    p.runs[0].font.name = 'Calibri'
    if level == 1: p.runs[0].text = text.upper()

def add_content_para_t1(doc, text):
    if text and text.strip().lower() != 'none':
        para_text = text.strip().replace('\n', ' ')
        if para_text:
            p = doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- Main Functions for Template 1 ---
def get_prompt_for_template_1(resume_text):
    """Creates the prompt for the Gemini API for Template 1."""
    template_instruction = """
You are a resume data extractor...
---
FullName: [Full Name]
Professional Summary:
[Extract and **summarize** the resume's professional overview...]
Roles:
[Extract all roles listed."]
Technologies:
[Extract the technologies... 'Category: Skills' format...]
Education:
[Extract content for the education section. **You MUST NOT extract any GPAs...**]
Certifications:
[Extract certifications...]
Geographic locale:
[Extract geographic locale...]
---JOB START---
CompanyName: [Company Name]
Role: [Your Role/Job Title]
Duration: [Start Date – End Date]
Client: [Client Name...]
Description: [Extract the project description]
Responsibilities:
- [Responsibility point 1]
---JOB END---
...
"""
    return f"Resume Text:\n{resume_text}\n\n{template_instruction}"

def parse_text_for_template_1(text):
    """Parses the tagged text from Gemini into a structured dictionary for Template 1."""
    resume_data = {"Jobs": []}
    lines = text.split('\n'); current_key = None
    main_keys = ["FullName", "Professional Summary", "Roles", "Technologies", "Education", "Certifications", "Geographic locale", "Professional and Experience Summary"]
    
    for line in lines:
        stripped_line = line.strip()
        if stripped_line == "---JOB START---":
            current_key = "Jobs"; resume_data["Jobs"].append({}); continue
        elif stripped_line == "---JOB END---":
            current_key = None; continue
        
        key_from_line, value_from_line = None, ""
        if ":" in line and not stripped_line.startswith('-'):
            try:
                key_from_line, value_from_line = line.split(":", 1)
                key_from_line, value_from_line = key_from_line.strip(), value_from_line.strip()
            except ValueError: pass 
        
        if key_from_line in main_keys:
            current_key = key_from_line; resume_data[current_key] = value_from_line; continue
        elif current_key == "Jobs" and key_from_line in ["CompanyName", "Role", "Duration", "Client", "BusinessValue", "Description", "Responsibilities"]:
            if key_from_line in ["CompanyName", "Role", "Duration", "Client", "BusinessValue", "Description"]:
                resume_data["Jobs"][-1][key_from_line] = value_from_line
            elif key_from_line == "Responsibilities":
                current_key = "Responsibilities" 
                if current_key not in resume_data["Jobs"][-1]: resume_data["Jobs"][-1][current_key] = []
                if value_from_line: resume_data["Jobs"][-1][current_key].append(value_from_line)
            continue
        elif current_key in main_keys and current_key != "FullName" and stripped_line:
            resume_data[current_key] += "\n" + line.strip()
        elif current_key == "Responsibilities" and stripped_line:
             if "Responsibilities" in resume_data["Jobs"][-1]:
                  resume_data["Jobs"][-1][current_key].append(stripped_line)
    return resume_data

def build_docx_for_template_1(resume_data):
    """Builds the DOCX document for Template 1."""
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    
    # --- Header (Logo Left, Name Right) ---
    header = doc.sections[0].header; header.is_linked_to_previous = False; header.paragraphs[0].text = "" 
    table_header = header.add_table(rows=1, cols=2, width=Inches(6.5)); set_table_no_border(table_header)
    cell_left = table_header.cell(0, 0); cell_left.width = Inches(1.5); p_left = cell_left.paragraphs[0]
    try:
        r_left = p_left.add_run(); r_left.add_picture('logo.png', width=Inches(1.5))
    except Exception: p_left.text = "[logo.png not found]"
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    cell_right = table_header.cell(0, 1); cell_right.width = Inches(5.0); p_right = cell_right.paragraphs[0]
    run_name = p_right.add_run(resume_data.get("FullName", "Candidate Name"))
    run_name.font.color.rgb = M_RED; run_name.font.name = 'Calibri'; run_name.bold = True; run_name.font.size = Pt(12)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_heading_t1(doc, "PROFESSIONAL OVERVIEW", level=1)
    add_content_para_t1(doc, resume_data.get("Professional Summary")); doc.add_paragraph() 
    
    overview_headings = ["Roles", "Technologies", "Education", "Certifications", "Geographic locale"]
    for heading in overview_headings:
        content_text = resume_data.get(heading, "None")
        if content_text.lower() == 'none' or not content_text.strip(): continue 

        p_heading = doc.add_paragraph(); run_heading = p_heading.add_run(heading + ":")
        run_heading.font.name = 'Calibri'; run_heading.bold = True; run_heading.font.color.rgb = M_RED

        if heading == "Roles" and content_text:
            # ... (T1 Roles logic) ...
            lines = content_text.strip().split('\n'); roles = []
            for line in lines: roles.extend([r.strip() for r in line.split(',') if r.strip()])
            for role in roles: doc.add_paragraph(role, style='List Bullet')
        elif heading == "Technologies" and content_text:
            # ... (T1 Technologies logic) ...
            tech_table = doc.add_table(rows=1, cols=2); tech_table.style = 'Table Grid' 
            tech_table.width = Inches(6.5); tech_table.columns[0].width = Inches(2.0); tech_table.columns[1].width = Inches(4.5)
            hdr_cells = tech_table.rows[0].cells; hdr_cells[0].paragraphs[0].add_run('Category').bold = True; hdr_cells[1].paragraphs[0].add_run('Skills').bold = True
            tech_groups = {}
            for line in content_text.strip().split('\n'):
                if ':' in line:
                    try:
                        category, skill = line.split(':', 1); category, skill = category.strip(), skill.strip()
                        if category in tech_groups: tech_groups[category].append(skill)
                        else: tech_groups[category] = [skill]
                    except ValueError: pass
            for category, skills_list in tech_groups.items():
                row_cells = tech_table.add_row().cells; row_cells[0].text = category; row_cells[1].text = ", ".join(skills_list)
            doc.add_paragraph() 
        elif heading == "Certifications" and content_text:
            # ... (T1 Certifications logic) ...
            lines = content_text.strip().split('\n')
            for line in lines:
                if line.strip(): doc.add_paragraph(line.lstrip('- '), style='List Bullet')
        else:
            for line in content_text.strip().split('\n'):
                if line.strip(): p = doc.add_paragraph(line.strip()); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph() 

    doc.add_page_break(); add_heading_t1(doc, "Professional and Experience Summary", level=1)
    for i, job_data in enumerate(resume_data.get("Jobs", [])):
        # ... (T1 Jobs formatting logic) ...
        add_heading_t1(doc, f"Project {i+1}", level=2)
        if job_data.get("Client"): p = doc.add_paragraph(); p.add_run("Client: ").bold = True; p.add_run(job_data.get("Client"))
        if job_data.get("Duration"): p = doc.add_paragraph(); p.add_run("Duration: ").bold = True; p.add_run(job_data.get("Duration"))
        if job_data.get("Role"): p = doc.add_paragraph(); p.add_run("Role: ").bold = True; p.add_run(job_data.get("Role"))
        if job_data.get("Description"): p = doc.add_paragraph(); p.add_run("Description: ").bold = True; add_content_para_t1(doc, job_data.get("Description"))
        responsibilities = job_data.get('Responsibilities', [])
        if responsibilities:
            p = doc.add_paragraph(); resp_run = p.add_run("Roles and Responsibilities:"); resp_run.font.name = 'Calibri'; resp_run.bold = True
            for resp in responsibilities:
                if resp.strip(): doc.add_paragraph(resp.lstrip('- '), style='List Bullet')
        doc.add_paragraph()

    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer


# --- 3. TEMPLATE 2 FUNCTIONS ---
# (This is the code you just provided)

# --- Helpers for Template 2 ---
def populate_table_cell_t2(cell, heading, content):
    p_heading = cell.paragraphs[0]; run_heading = p_heading.add_run(heading)
    run_heading.font.name = 'Lato'; run_heading.bold = True; run_heading.font.size = Pt(10)
    
    # Split content by newline for bullet points or multiple lines
    content_lines = content.strip().split('\n')
    for i, line in enumerate(content_lines):
        if line.strip(): # Avoid adding empty paragraphs
            # Use add_paragraph for all lines to ensure consistent styling
            p_content = cell.add_paragraph(line.strip())
            p_content.style.font.name = 'Lato'
            p_content.style.font.size = Pt(11)
            # Remove the default paragraph if it's empty (created by cell)
            if i == 0 and not cell.paragraphs[0].text:
                 p_heading.clear() # Clear if we are replacing it
                 p_heading._element.getparent().remove(p_heading._element)

# --- Main Functions for Template 2 ---
def get_prompt_for_template_2(resume_text):
    """Creates the prompt for the Gemini API for Template 2."""
    template_instruction = """
You are a resume data extractor. Your task is to extract information from the provided resume and format it as clean, tagged, plain text.

DO NOT add any special formatting. The Python script will handle all styling.

---

FullName: [Full Name]
Designation: [Designation]

ProfessionalOverviewSummary:
[A 2-3 sentence summary of the professional profile, extracted from the resume.Generate based on resume if not explicitly mentioned]

ProfessionalOverviewTable:
Roles | [Summarize key roles held, comma separated]
Solutions | [Summarize solutions delivered]
Industries | [List relevant industries]
Technologies | [List key technologies used]

KeyEngagementsTable:
Client | Role | Description
[Client Name 1] | [Role at Client 1] | [Brief description of engagement 1]
[Client Name 2] | [Role at Client 2] | [Brief description of engagement 2]

Education:
[Content for the education section]

Publications:
[Content for the publications section]

ProfessionalTrainingCertifications:
[Content for certifications section]

GeographicLocale:
[Content for geographic locale section]


---JOB START---
CompanyName: [Company Name]
Role: [Your Role/Job Title]
Duration: [Start Date – End Date]
Client: [Client Name for the project. If not applicable, write N/A]
Responsibilities:
- [Responsibility point 1]
---JOB END---

Repeat the ---JOB START--- to ---JOB END--- block for each job. If a section is empty, write "None".
"""
    return f"Resume Text:\n{resume_text}\n\n{template_instruction}"

def parse_text_for_template_2(text):
    """Parses the tagged text from Gemini into a structured dictionary for Template 2."""
    resume_data = {}
    lines = text.split('\n'); current_key = None
    
    for line in lines:
        stripped_line = line.strip()

        # Handle JOB markers first
        if stripped_line == "---JOB START---":
            current_key = "Jobs"
            if current_key not in resume_data: resume_data[current_key] = []
            resume_data[current_key].append({})
            continue
        elif stripped_line == "---JOB END---":
            current_key = None
            continue

        # Handle lines with colons
        if ":" in line and not stripped_line.startswith('-'):
            try:
                key, value = line.split(":", 1)
                key = key.strip()

                if key in ["ProfessionalOverviewSummary", "Education", "Publications", "ProfessionalTrainingCertifications", "GeographicLocale", "ProfessionalOverviewTable", "KeyEngagementsTable"]:
                    current_key = key
                    resume_data[current_key] = value.strip() + "\n" # Start multi-line capture
                elif current_key == "Jobs" and resume_data.get("Jobs"):
                    if key in ["CompanyName", "Role", "Duration", "Client"]:
                        resume_data["Jobs"][-1][key] = value.strip()
                    elif key == "Responsibilities":
                        current_key = "Responsibilities"
                        if current_key not in resume_data["Jobs"][-1]:
                            resume_data["Jobs"][-1][current_key] = []
                        # Handle if a responsibility is on the same line
                        if value.strip():
                            resume_data["Jobs"][-1][current_key].append(value.strip().lstrip('- '))
                else: # Simple key-value pairs like FullName
                    resume_data[key] = value.strip()
                    current_key = None # Reset current key
            except ValueError:
                pass # Line might have a colon but not be a key-value pair
        
        # Handle multi-line content and bullet points
        elif current_key:
            if current_key == "Responsibilities" and resume_data.get("Jobs"):
                if stripped_line: # Only add non-empty lines
                    resume_data["Jobs"][-1][current_key].append(stripped_line.lstrip('- '))
            elif current_key in resume_data:
                resume_data[current_key] += line + "\n"
    return resume_data

def build_docx_for_template_2(resume_data):
    """Builds the DOCX document for Template 2."""
    doc = Document()
    style = doc.styles['Normal']; font = style.font
    font.name = 'Lato'; font.size = Pt(11)

    # --- Header with Wave (header.png) and Logo (logo.png) ---
    try:
        header = doc.sections[0].header
        for paragraph in list(header.paragraphs): header._element.remove(paragraph._element)
        
        p_wave = header.add_paragraph()
        p_wave.paragraph_format.space_before = Pt(0)
        p_wave.paragraph_format.space_after = Pt(0)
        run_wave = p_wave.add_run()
        run_wave.add_picture('header.png', width=Inches(8.5)) 

        p_logo = header.add_paragraph()
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(0)
        run_logo = p_logo.add_run()
        run_logo.add_picture('logo.png', width=Inches(1.5))
    except Exception as e:
        print(f"Info: 'header.png' or 'logo.png' not found. Error: {e}")

    # --- Add footer.png border to footer ---
    try:
        footer = doc.sections[0].footer
        for paragraph in list(footer.paragraphs): footer._element.remove(paragraph._element)
            
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = footer_para.add_run()
        run_footer.add_picture('footer.png', width=Inches(8.5)) 
    except Exception as e:
        print(f"Info: 'footer.png' not found. Error: {e}")

    # --- Add Name and Designation to FIRST PAGE BODY ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_name = p_name.add_run(resume_data.get("FullName", "Candidates Name"))
    run_name.font.color.rgb = RGBColor(204, 31, 32) # Red
    run_name.font.name = "Lato"; run_name.bold = True; run_name.font.size = Pt(18)

    p_des = doc.add_paragraph()
    p_des.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_des = p_des.add_run(resume_data.get("Designation", "Designation(Latest)"))
    run_des.font.color.rgb = RGBColor(0, 0, 0) # Black
    run_des.font.name = "Lato"; run_des.bold = False; run_des.font.size = Pt(12)
    doc.add_paragraph() # Spacer

    doc.add_paragraph("Professional Overview:", style='Heading 2').runs[0].font.color.rgb = RGBColor(204, 31, 32)
    doc.add_paragraph(resume_data.get("ProfessionalOverviewSummary", "").strip())
    
    # --- Professional Overview Table with Red Bullets for Roles ---
    table_lines_po = resume_data.get("ProfessionalOverviewTable", "").strip().split('\n')
    if table_lines_po and '|' in table_lines_po[0]:
        table_data = []
        for row_str in table_lines_po:
            if '|' in row_str:
                parts = [cell.strip() for cell in row_str.split('|', 1)]
                if len(parts) == 2: table_data.append(parts)

        if table_data:
            table = doc.add_table(rows=len(table_data), cols=2)
            table.style = 'Table Grid'
            
            for r, row_data in enumerate(table_data):
                heading, content = row_data
                
                heading_cell = table.cell(r, 0)
                heading_cell.text = heading
                heading_cell.paragraphs[0].runs[0].font.bold = True
                heading_cell.paragraphs[0].runs[0].font.name = 'Lato'

                content_cell = table.cell(r, 1)
                content_cell.text = "" # Clear default

                if heading.lower() == "roles":
                    roles_list = [role.strip() for role in content.split(',') if role.strip()]
                    for i, role in enumerate(roles_list):
                        p_bullet = content_cell.paragraphs[0] if i == 0 else content_cell.add_paragraph()
                        p_bullet.paragraph_format.space_before = Pt(0)
                        p_bullet.paragraph_format.space_after = Pt(0)

                        run_bullet = p_bullet.add_run('•') # Bullet
                        run_bullet.font.color.rgb = RGBColor(204, 31, 32) # Red
                        run_bullet.font.name = 'Lato'
                        p_bullet.add_run('\t') # Tab
                        run_text = p_bullet.add_run(role) # Role text
                        run_text.font.name = 'Lato'; run_text.font.size = Pt(11)
                        
                        p_bullet.paragraph_format.left_indent = Inches(0.25)
                        p_bullet.paragraph_format.first_line_indent = Inches(-0.25)
                else:
                    content_cell.text = content
                    content_cell.paragraphs[0].runs[0].font.name = 'Lato'
                    content_cell.paragraphs[0].runs[0].font.size = Pt(11)
            doc.add_paragraph()

    doc.add_paragraph().add_run("Key Engagements").italic = True
    table_lines_ke = resume_data.get("KeyEngagementsTable", "").strip().split('\n')
    if table_lines_ke and '|' in table_lines_ke[0]:
        table_data = []
        for row in table_lines_ke:
             parts = [cell.strip() for cell in row.split('|')]
             if len(parts) > 1: table_data.append(parts) # Ensure valid row
        
        if table_data:
            num_cols = len(table_data[0])
            table = doc.add_table(rows=len(table_data), cols=num_cols); table.style = 'Table Grid'
            for r, row_data in enumerate(table_data):
                if len(row_data) == num_cols: # Only process rows that match header
                    for c, cell_data in enumerate(row_data): 
                        table.cell(r, c).text = cell_data
                        # Style header row
                        if r == 0:
                            table.cell(r,c).paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

    # --- 2x2 Table ---
    table_2x2 = doc.add_table(rows=2, cols=2); set_table_no_border(table_2x2)
    # Re-using T2's helper, but it needs to be modified to handle multi-line
    # Let's adjust the populate_table_cell_t2 helper and how it's called
    
    # Simple T2 implementation:
    def populate_cell_t2_simple(cell, heading, content):
        p_heading = cell.paragraphs[0]; run_heading = p_heading.add_run(heading)
        run_heading.font.name = 'Lato'; run_heading.bold = True; run_heading.font.size = Pt(10)
        
        content_lines = content.strip().split('\n')
        for line in content_lines:
            if line.strip():
                p_content = cell.add_paragraph(line.strip())
                p_content.style.font.name = 'Lato'
                p_content.style.font.size = Pt(11)

    populate_cell_t2_simple(table_2x2.cell(0, 0), "Education", resume_data.get("Education", "None").strip())
    populate_cell_t2_simple(table_2x2.cell(0, 1), "Professional Training/Certifications", resume_data.get("ProfessionalTrainingCertifications", "None").strip())
    populate_cell_t2_simple(table_2x2.cell(1, 0), "Publications", resume_data.get("Publications", "None").strip())
    populate_cell_t2_simple(table_2x2.cell(1, 1), "Geographic locale", resume_data.get("GeographicLocale", "None").strip())
    doc.add_paragraph()


    doc.add_page_break()
    doc.add_paragraph("Professional and Business Experience:", style='Heading 2').runs[0].font.color.rgb = RGBColor(204, 31, 32)

    for job_data in resume_data.get("Jobs", []):
        p = doc.add_paragraph(); company_run = p.add_run(job_data.get("CompanyName", "")); company_run.font.color.rgb = RGBColor(204, 31, 32); company_run.font.name = 'Lato'; company_run.bold = True
        p.add_run('\t'); duration_run = p.add_run(job_data.get("Duration", "")); duration_run.font.name = 'Lato'; p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        p = doc.add_paragraph(); role_run = p.add_run(job_data.get("Role", "")); role_run.font.name = 'Lato'; role_run.bold = True; doc.add_paragraph()
        p = doc.add_paragraph(); client_label_run = p.add_run("CLIENT: "); client_label_run.font.name = 'Lato'; client_label_run.bold = True; client_text_run = p.add_run(job_data.get("Client", "N/A")); client_text_run.font.name = 'Lato'; doc.add_paragraph()
        p = doc.add_paragraph(); resp_run = p.add_run("Responsibilities:"); resp_run.font.name = 'Lato'; resp_run.underline = True
        
        for resp in job_data.get('Responsibilities', []):
            if resp.strip(): # Avoid empty bullet points
                doc.add_paragraph(resp.strip(), style='List Bullet')
        doc.add_paragraph()
        
    
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)

    return buffer
